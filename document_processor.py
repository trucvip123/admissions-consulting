import os
from docx import Document
from typing import List, Dict
import logging
from config import Config

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentProcessor:
    def __init__(self, data_dir: str = Config.DATA_DIR):
        self.data_dir = data_dir

    def extract_text_from_docx(self, file_path: str) -> str:
        """Trích xuất văn bản từ file docx"""
        try:
            doc = Document(file_path)
            text = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text.strip())

            # Trích xuất từ bảng
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text.append(" | ".join(row_text))

            return "\n".join(text)
        except Exception as e:
            logger.error(f"Lỗi khi đọc file {file_path}: {str(e)}")
            return ""

    def process_all_documents(self) -> List[Dict]:
        """Xử lý tất cả tài liệu trong thư mục data"""
        documents = []

        for filename in os.listdir(self.data_dir):
            if filename.endswith(".docx"):
                file_path = os.path.join(self.data_dir, filename)
                logger.info(f"Đang xử lý file: {filename}")

                content = self.extract_text_from_docx(file_path)
                if content:
                    documents.append(
                        {"filename": filename, "content": content, "source": file_path}
                    )
                    logger.info(
                        f"Đã xử lý thành công: {filename} ({len(content)} ký tự)"
                    )

        return documents

    def get_document_summary(self) -> Dict:
        """Tạo tóm tắt về các tài liệu đã xử lý"""
        documents = self.process_all_documents()

        summary = {
            "total_documents": len(documents),
            "total_characters": sum(len(doc["content"]) for doc in documents),
            "documents": [],
        }

        for doc in documents:
            summary["documents"].append(
                {
                    "filename": doc["filename"],
                    "characters": len(doc["content"]),
                    "lines": len(doc["content"].split("\n")),
                }
            )

        return summary


if __name__ == "__main__":
    processor = DocumentProcessor()
    summary = processor.get_document_summary()
    print("Tóm tắt tài liệu:")
    print(f"Tổng số tài liệu: {summary['total_documents']}")
    print(f"Tổng số ký tự: {summary['total_characters']}")
    for doc in summary["documents"]:
        print(f"- {doc['filename']}: {doc['characters']} ký tự, {doc['lines']} dòng")
